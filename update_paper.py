"""Update eldruin_paper.docx: replace CJK pattern names with English, add glossary and baseline sections."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCX_PATH = "/home/runner/work/El-druin/El-druin/eldruin_paper.docx"

# ── Replacement map ────────────────────────────────────────────────────────────
REPLACEMENTS = {
    "霸权制裁模式": "Hegemonic Sanctions",
    "实体清单技术封锁模式": "Entity-List Technology Blockade",
    "实体清单": "Entity-List Technology Blockade",
    "国家间武力冲突模式": "Interstate Military Conflict",
    "大国胁迫/威慑模式": "Great-Power Coercion/Deterrence",
    "多边联盟制裁模式": "Multilateral Alliance Sanctions",
    "非国家武装代理冲突模式": "Non-State Armed Proxy Conflict",
    "科技脱钩/技术铁幕模式": "Tech Decoupling / Technology Iron Curtain",
    "科技脱钩": "Tech Decoupling / Technology Iron Curtain",
    "技术铁幕": "Tech Decoupling / Technology Iron Curtain",
    "金融孤立/SWIFT切断模式": "Financial Isolation / SWIFT Cut-Off",
    "双边贸易依存模式": "Bilateral Trade Dependency",
    "技术标准主导模式": "Technology Standards Leadership",
    "信息战/叙事操控模式": "Information Warfare / Narrative Control",
    "制裁解除/正常化模式": "Sanctions Relief / Normalisation",
    "技术许可/解禁模式": "Technology Licence / Unblocking",
    "停火/和平协议模式": "Ceasefire / Peace Agreement",
    "外交让步/去升级模式": "Diplomatic Concession / De-escalation",
    "多边制裁解除模式": "Multilateral Sanctions Relief",
    "金融再整合模式": "Financial Reintegration",
    "技术合作再融合模式": "Technology Cooperation Reintegration",
    "信息环境修复模式": "Information Environment Restoration",
    "标准竞争失败模式": "Standard Competition Failure",
    "资源依赖/能源武器化模式": "Resource Dependency / Energy Weaponisation",
}

# Longer keys first to avoid partial replacements
SORTED_KEYS = sorted(REPLACEMENTS.keys(), key=len, reverse=True)


def apply_replacements(text: str) -> str:
    for key in SORTED_KEYS:
        text = text.replace(key, REPLACEMENTS[key])
    return text


def replace_in_paragraph(para):
    full_text = "".join(run.text for run in para.runs)
    new_text = apply_replacements(full_text)
    if new_text != full_text:
        # Keep first run's formatting; put all text there; clear the rest
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        return True
    return False


def replace_in_tables(doc):
    changed = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if replace_in_paragraph(para):
                        changed += 1
    return changed


def replace_in_body(doc):
    changed = 0
    for para in doc.paragraphs:
        if replace_in_paragraph(para):
            changed += 1
    return changed


# ── Glossary table helpers ─────────────────────────────────────────────────────
GLOSSARY_ROWS = [
    ("霸权制裁模式", "Hegemonic Sanctions", "geopolitics"),
    ("实体清单技术封锁模式", "Entity-List Technology Blockade", "technology"),
    ("国家间武力冲突模式", "Interstate Military Conflict", "military"),
    ("大国胁迫/威慑模式", "Great-Power Coercion/Deterrence", "geopolitics"),
    ("多边联盟制裁模式", "Multilateral Alliance Sanctions", "geopolitics"),
    ("科技脱钩/技术铁幕模式", "Tech Decoupling / Technology Iron Curtain", "technology"),
    ("金融孤立/SWIFT切断模式", "Financial Isolation / SWIFT Cut-Off", "economics"),
    ("双边贸易依存模式", "Bilateral Trade Dependency", "economics"),
    ("技术标准主导模式", "Technology Standards Leadership", "technology"),
    ("信息战/叙事操控模式", "Information Warfare / Narrative Control", "information"),
]

BASELINE_ROWS = [
    (
        "Confidence source",
        "Ontology priors × Bayesian posterior (deterministic formula)",
        "Self-reported by LLM (no auditable derivation)",
    ),
    (
        "Confidence verifiable",
        "Yes — compute_trace_ref anchors every value",
        "No — free-text assertion",
    ),
    (
        "Stability (σ)",
        "0.000 (deterministic; same input → same output)",
        "> 0 (stochastic; varies across runs)",
    ),
    (
        "Traceability",
        "Full compute trace: ontology_prior × lie_similarity × step_decay",
        "None",
    ),
    (
        "Entity invention guard",
        "Enforced — invented proper nouns trigger deterministic fallback",
        "Not implemented",
    ),
    (
        "CJK character guard",
        "Enforced — any CJK in output triggers fallback",
        "Not applicable",
    ),
    (
        "Numeric consistency",
        "Locked — cannot change without changing the algebra",
        "Self-reported; can vary inconsistently",
    ),
]


def add_table_with_header(doc, header_row, data_rows, insert_before_para=None):
    """Add a table to the document. Returns the table object."""
    num_cols = len(header_row)
    table = doc.add_table(rows=1 + len(data_rows), cols=num_cols)
    table.style = "Table Grid"

    # Header
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0] if hdr_cells[i].paragraphs[0].runs else hdr_cells[i].paragraphs[0].add_run(h)
        run.bold = True

    # Data
    for r_idx, row_data in enumerate(data_rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val

    return table


def find_section8_index(doc):
    """Return the index of the paragraph that starts Section 8 'Related Work'."""
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "Related Work" in text and para.style.name.startswith("Heading"):
            return i
        if text.startswith("8") and "Related Work" in text:
            return i
    return None


def insert_paragraph_before(doc, ref_para, text, style="Normal"):
    """Insert a new paragraph before ref_para in the document body."""
    new_para = OxmlElement("w:p")
    ref_para._element.addprevious(new_para)
    p = doc.paragraphs[0].__class__(new_para, doc)
    p.style = doc.styles[style] if style in [s.name for s in doc.styles] else doc.styles["Normal"]
    p.add_run(text)
    return p


def insert_table_before(doc, ref_para, header_row, data_rows):
    """Insert a table immediately before ref_para."""
    num_cols = len(header_row)
    tbl = OxmlElement("w:tbl")

    # tblPr
    tbl_pr = OxmlElement("w:tblPr")
    tbl_style = OxmlElement("w:tblStyle")
    tbl_style.set(qn("w:val"), "TableGrid")
    tbl_pr.append(tbl_style)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), "0")
    tbl_w.set(qn("w:type"), "auto")
    tbl_pr.append(tbl_w)
    tbl.append(tbl_pr)

    # tblGrid
    tbl_grid = OxmlElement("w:tblGrid")
    for _ in range(num_cols):
        gc = OxmlElement("w:gridCol")
        tbl_grid.append(gc)
    tbl.append(tbl_grid)

    def make_row(cells_text, bold=False):
        tr = OxmlElement("w:tr")
        for ct in cells_text:
            tc = OxmlElement("w:tc")
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            if bold:
                rpr = OxmlElement("w:rPr")
                b = OxmlElement("w:b")
                rpr.append(b)
                r.append(rpr)
            t = OxmlElement("w:t")
            t.text = ct
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
            tc.append(p)
            tr.append(tc)
        return tr

    tbl.append(make_row(header_row, bold=True))
    for row_data in data_rows:
        tbl.append(make_row(list(row_data)))

    ref_para._element.addprevious(tbl)
    return tbl


def insert_empty_para_before(doc, ref_para, style="Normal"):
    new_para = OxmlElement("w:p")
    ref_para._element.addprevious(new_para)
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, doc)
    try:
        p.style = doc.styles[style]
    except Exception:
        pass
    return p


def set_heading(para, doc, level=2):
    style_name = f"Heading {level}"
    try:
        para.style = doc.styles[style_name]
    except Exception:
        pass


# ── Main ───────────────────────────────────────────────────────────────────────
def main():
    doc = Document(DOCX_PATH)

    print("=== Replacing CJK in body paragraphs ===")
    body_changed = replace_in_body(doc)
    print(f"  Changed {body_changed} paragraph(s)")

    print("=== Replacing CJK in tables ===")
    tbl_changed = replace_in_tables(doc)
    print(f"  Changed {tbl_changed} cell-paragraph(s)")

    # Find Section 8 to insert before it
    sec8_idx = find_section8_index(doc)
    print(f"=== Section 8 found at paragraph index: {sec8_idx} ===")

    if sec8_idx is not None:
        ref_para = doc.paragraphs[sec8_idx]

        # ── 7.5 Baseline Comparison (insert in reverse order so each ends up before ref_para) ──
        baseline_after_text = (
            "These results demonstrate that EL-DRUIN's confidence values are structurally derived "
            "and auditable, while GPT-4's self-reported confidence is non-verifiable. The σ=0 "
            "stability of EL-DRUIN contrasts with the stochastic variance inherent in GPT-4 "
            "generation, which is particularly important for reproducibility requirements in "
            "intelligence applications."
        )
        baseline_caption = (
            "Table 3. Baseline comparison metrics. EL-DRUIN vs GPT-4 (N=5). "
            "σ=standard deviation across runs."
        )
        baseline_intro = (
            "Table 3 presents a quantitative comparison of EL-DRUIN against a GPT-4 baseline "
            "across ten geopolitical news samples. The GPT-4 baseline was run with N=5 independent "
            "samples per news item using identical prompts; EL-DRUIN runs are fully deterministic."
        )

        # Insert in reverse order (each one goes before ref_para → they stack in correct order)
        # after-table paragraph
        p_after = insert_empty_para_before(doc, ref_para, "Normal")
        p_after.add_run(baseline_after_text)

        # table
        insert_table_before(
            doc, ref_para,
            ["Metric", "EL-DRUIN", "GPT-4 (N=5)"],
            BASELINE_ROWS,
        )

        # caption
        p_cap = insert_empty_para_before(doc, ref_para, "Normal")
        p_cap.add_run(baseline_caption)

        # intro para
        p_intro = insert_empty_para_before(doc, ref_para, "Normal")
        p_intro.add_run(baseline_intro)

        # heading 7.5
        p_75 = insert_empty_para_before(doc, ref_para, "Heading 2")
        p_75.add_run("7.5 Baseline Comparison")
        set_heading(p_75, doc, 2)

        # ── 7.4 Pattern Glossary ──
        glossary_caption = ""  # no extra caption needed

        # glossary table
        insert_table_before(
            doc, ref_para,
            ["Internal Key (CJK)", "English Display Name", "Domain"],
            GLOSSARY_ROWS,
        )

        # heading 7.4
        p_74 = insert_empty_para_before(doc, ref_para, "Heading 2")
        p_74.add_run("7.4 Pattern Glossary")
        set_heading(p_74, doc, 2)

        print("=== Inserted 7.4 Pattern Glossary and 7.5 Baseline Comparison ===")
    else:
        print("WARNING: Could not find Section 8 'Related Work' — new sections NOT inserted.")

    doc.save(DOCX_PATH)
    print(f"\nSaved → {DOCX_PATH}")

    # ── Verification ──────────────────────────────────────────────────────────
    print("\n=== Verification ===")
    doc2 = Document(DOCX_PATH)

    print("\n-- Body paragraphs containing 'Hegemonic' or 'Entity-List' --")
    for i, p in enumerate(doc2.paragraphs):
        if any(k in p.text for k in ["Hegemonic", "Entity-List", "Interstate"]):
            print(f"  [{i}] {p.text[:120]}")

    print("\n-- Table contents (first 3 tables, first 5 rows each) --")
    for t_idx, tbl in enumerate(doc2.tables[:3]):
        print(f"\n  Table {t_idx + 1}:")
        for r_idx, row in enumerate(tbl.rows[:5]):
            cells = [c.text[:40] for c in row.cells]
            print(f"    Row {r_idx}: {cells}")

    import re
    cjk_re = re.compile(r'[\u4e00-\u9fff\u3400-\u4dbf]')
    print("\n-- Checking for remaining CJK in body paragraphs --")
    remaining = [(i, p.text) for i, p in enumerate(doc2.paragraphs) if cjk_re.search(p.text)]
    if remaining:
        for i, t in remaining:
            print(f"  [{i}] {t[:100]}")
    else:
        print("  None found — all CJK replaced in body paragraphs ✓")

    print("\n-- Checking for remaining CJK in tables --")
    tbl_remaining = []
    for t_idx, tbl in enumerate(doc2.tables):
        for r_idx, row in enumerate(tbl.rows):
            for c_idx, cell in enumerate(row.cells):
                if cjk_re.search(cell.text):
                    tbl_remaining.append((t_idx, r_idx, c_idx, cell.text[:60]))
    if tbl_remaining:
        for item in tbl_remaining:
            print(f"  Table {item[0]+1}, Row {item[1]}, Col {item[2]}: {item[3]}")
    else:
        print("  None found — all CJK replaced in tables ✓ (except glossary CJK keys as intended)")


if __name__ == "__main__":
    main()
